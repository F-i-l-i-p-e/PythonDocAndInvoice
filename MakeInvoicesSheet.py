from docx import Document
import openpyxl
import os
import re

# Function to extract invoice data from a docx file
def extract_invoice_data(doc):
    invoice_num = "INV" + doc.paragraphs[0].text[3:]
    products_purchased = 0
    subtotal, tax, total = 0, 0, 0

    for paragraph in doc.paragraphs[1:]:
        if "PRODUCTS" in paragraph.text:
            products_purchased = sum([int(x.split(':')[1]) for x in re.findall(r'\w+:\d+', paragraph.text)])
        elif "SUBTOTAL" in paragraph.text:
            values = re.findall(r'\d+\.\d+', paragraph.text)
            subtotal, tax, total = float(values[0]), float(values[1]), float(values[2])

    return (invoice_num, products_purchased, subtotal, tax, total)

# Create sheet
wb = openpyxl.Workbook()
ws = wb.active
ws.title = "Invoice Report"

# Write column headers
column_headers = ["Invoice Number", "Total Quantity", "Subtotal", "Tax", "Total"]
for col, header in enumerate(column_headers, start=1):
    ws.cell(row=1, column=col, value=header)

# Read docx files and send data to sheet
current_row = 2
for file in os.listdir():
    if file.endswith('.docx'):
        with open(file, 'rb') as file_handle:
            doc = Document(file_handle)
            invoice_data = extract_invoice_data(doc)
            ws.append(invoice_data)
            current_row += 1

# Make header bold
for col in range(1, 6):
    ws.cell(row=1, column=col).font = openpyxl.styles.Font(bold=True)

# Save spreadsheet
output_file = "Invoices_Sheet.xlsx"
wb.save(output_file)

